from fastapi import (
    FastAPI,
    UploadFile,
    File,
    HTTPException,
    Form,
    Response,
    Depends,
    Query,
)
from fastapi.responses import FileResponse, JSONResponse
from pathlib import Path
import tempfile
from sqlalchemy import create_engine, Column, Integer, String
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from pydantic import BaseModel
from typing import List, Optional, Literal
from so_vits_svc_fork.inference.main import infer
import os
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.exc import NoResultFound
from sqlalchemy.future import select
import uuid
from gtts import gTTS
from docx import Document
import io
import logging
from logging import getLogger
from so_vits_svc_fork.preprocessing.preprocess_split import preprocess_split
from so_vits_svc_fork.preprocessing.preprocess_resample import preprocess_resample
from so_vits_svc_fork.preprocessing.preprocess_flist_config import (
    CONFIG_TEMPLATE_DIR,
    preprocess_config,
)
from so_vits_svc_fork.preprocessing.preprocess_hubert_f0 import preprocess_hubert_f0
from so_vits_svc_fork.train import train
import shutil
import shutil
import asyncio
import aiofiles
import uvicorn

# Kết nối tới cơ sở dữ liệu MySQL
DATABASE_URL = "mysql+pymysql://root:g08022002@localhost/datn"  # Thay đổi username, password, và dbname của bạn
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

BASE_DIR = Path(__file__).resolve().parent
section_storage_path = BASE_DIR / "files"
train_model_path = BASE_DIR / "trainmodel"


# Định nghĩa cơ sở dữ liệu
Base = declarative_base()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class Model(Base):
    __tablename__ = "models"  # Tên bảng trong MySQL

    id_model = Column(Integer, primary_key=True, index=True)
    model_name = Column(String(100), index=True)
    model_path = Column(String(200))
    config_path = Column(String(200))
    cluster_model_path = Column(String(200), nullable=True)
    category = Column(String(200), nullable=True)


# Tạo FastAPI app
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "*"
    ],  # Cho phép tất cả các nguồn. Bạn có thể điều chỉnh để chỉ cho phép các nguồn cụ thể.
    allow_credentials=True,
    allow_methods=[
        "*"
    ],  # Cho phép tất cả các phương thức HTTP (GET, POST, PUT, DELETE, v.v.).
    allow_headers=["*"],  # Cho phép tất cả các header.
)


# Định nghĩa schema cho phản hồi
class ModelResponse(BaseModel):
    id_model: int
    model_name: str
    model_path: str
    config_path: str
    cluster_model_path: str
    category: str

    class Config:
        from_attributes = True
        protected_namespaces = ()


class TextToSpeechRequest(BaseModel):
    text: str
    locate: str = "vi"


class TextToSpeechAndInferRequest(BaseModel):
    text: str
    locate: str = "vi"
    model_id: int


# Tạo một phiên làm việc với cơ sở dữ liệu
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


async def save_uploaded_file(uploaded_file, destination: Path):
    try:
        with destination.open("wb") as dest_file:
            while True:
                chunk = await asyncio.to_thread(
                    uploaded_file.read, 1024
                )  # Đọc từng khối 1024 byte
                if not chunk:
                    break
                dest_file.write(chunk)
    except Exception as e:
        raise RuntimeError(f"Error saving uploaded file: {e}")


def pre_split(
    input_dir: str,
    output_dir: str,
    sr: int,
    max_length: float = 5.0,
    top_db: int = 30,
    frame_seconds: float = 0.5,
    hop_seconds: float = 0.1,
    n_jobs: int = -1,
) -> None:

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    preprocess_split(
        input_dir=input_dir,
        output_dir=output_dir,
        sr=sr,
        max_length=max_length,
        top_db=top_db,
        frame_seconds=frame_seconds,
        hop_seconds=hop_seconds,
        n_jobs=n_jobs,
    )


def process_audio_files(
    input_dir: str,
    output_dir: str,
    sampling_rate: int = 16000,
    top_db: int = 30,
    frame_seconds: float = 0.1,
    hop_seconds: float = 0.05,
    n_jobs: int = -1,
) -> None:
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    # Tạo thư mục output nếu chưa tồn tại
    output_path.mkdir(parents=True, exist_ok=True)

    # Gọi hàm preprocess_resample để xử lý các tệp âm thanh
    preprocess_resample(
        input_dir=input_path,
        output_dir=output_path,
        sampling_rate=sampling_rate,
        n_jobs=n_jobs,
        top_db=top_db,
        frame_seconds=frame_seconds,
        hop_seconds=hop_seconds,
    )


def pre_config(
    input_dir: Path,
    filelist_path: Path,
    config_path: Path,
    config_type: str,
):

    input_dir = Path(input_dir)
    filelist_path = Path(filelist_path)
    config_path = Path(config_path)
    preprocess_config(
        input_dir=input_dir,
        train_list_path=filelist_path / "train.txt",
        val_list_path=filelist_path / "val.txt",
        test_list_path=filelist_path / "test.txt",
        config_path=config_path,
        config_name=config_type,
    )

    print("Configuration processing completed.")


def pre_hubert(
    input_dir: Path,
    config_path: Path,
    n_jobs: bool,
    force_rebuild: bool,
    f0_method: Literal["crepe", "crepe-tiny", "parselmouth", "dio", "harvest"],
) -> None:
    from so_vits_svc_fork.preprocessing.preprocess_hubert_f0 import preprocess_hubert_f0

    input_dir = Path(input_dir)
    config_path = Path(config_path)
    preprocess_hubert_f0(
        input_dir=input_dir,
        config_path=config_path,
        n_jobs=n_jobs,
        force_rebuild=force_rebuild,
        f0_method=f0_method,
    )


def train(
    config_path: Path,
    model_path: Path,
    tensorboard: bool = False,
    reset_optimizer: bool = False,
):
    config_path = Path(config_path)
    model_path = Path(model_path)

    if tensorboard:
        import webbrowser
        from tensorboard import program

        getLogger("tensorboard").setLevel(30)
        tb = program.TensorBoard()
        tb.configure(argv=[None, "--logdir", model_path.as_posix()])
        url = tb.launch()
        webbrowser.open(url)

    # Gọi hàm huấn luyện thực tế từ module khác, nếu cần
    from so_vits_svc_fork.train import train  # Chỉ gọi một lần từ đây

    train(
        config_path=config_path, model_path=model_path, reset_optimizer=reset_optimizer
    )


# Route để lấy danh sách các mô hình từ MySQL với khả năng lọc theo category
@app.get("/models", response_model=List[ModelResponse])
async def get_models(
    db: Session = Depends(get_db),
    category: Optional[str] = Query(None, title="Category to filter"),
):
    if category:
        # Lọc mô hình theo category
        models = db.query(Model).filter(Model.category == category).all()
    else:
        # Lấy tất cả mô hình nếu không có filter
        models = db.query(Model).all()

    return models


@app.post("/models/", response_model=ModelResponse)
async def create_model(
    model_name: str = Form(...),
    model_path: str = Form(...),
    config_path: str = Form(...),
    cluster_model_path: str = Form(None),
    category: str = Form(...),
    db: Session = Depends(get_db),
):
    new_model = Model(
        model_name=model_name,
        model_path=model_path,
        config_path=config_path,
        cluster_model_path=cluster_model_path,
        category=category,
    )
    db.add(new_model)
    db.commit()
    db.refresh(new_model)
    return new_model


@app.put("/models/{model_id}", response_model=ModelResponse)
async def update_model(
    model_id: int,
    model_name: str = Form(...),
    model_path: str = Form(...),
    config_path: str = Form(...),
    cluster_model_path: str = Form(None),
    category: str = Form(...),
    db: Session = Depends(get_db),
):
    model_info = db.query(Model).filter(Model.id_model == model_id).first()
    if not model_info:
        raise HTTPException(status_code=404, detail="Model not found")

    model_info.model_name = model_name
    model_info.model_path = model_path
    model_info.config_path = config_path
    model_info.cluster_model_path = cluster_model_path
    model_info.category = category

    db.commit()
    db.refresh(model_info)
    return model_info


@app.delete("/models/{model_id}", response_model=dict)
async def delete_model(model_id: int, db: Session = Depends(get_db)):
    model_info = db.query(Model).filter(Model.id_model == model_id).first()
    if not model_info:
        raise HTTPException(status_code=404, detail="Model not found")

    db.delete(model_info)
    db.commit()
    return {"detail": "Model deleted successfully"}


# API TTS
@app.post("/text-to-speech/")
async def text_to_speech(request: TextToSpeechRequest):
    section_id = str(uuid.uuid4())
    audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

    # Generate audio file
    try:
        tts = gTTS(text=request.text, lang=request.locate)
        # Lưu file âm thanh trong một luồng riêng biệt
        await asyncio.to_thread(tts.save, section_cloned_file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating audio: {str(e)}")

    return FileResponse(audio_file_path, media_type="audio/wav", filename="output.wav")


@app.post("/text-file-to-speech/")
async def text_file_to_speech(
    file: UploadFile = File(...),
    locate: str = Form("vi"),
):
    if file.filename.endswith(".txt"):
        # Đọc nội dung từ tệp văn bản
        content = await file.read()
        text = content.decode("utf-8")  # Giải mã nội dung
        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

        # Chuyển đổi văn bản thành giọng nói
        try:
            tts = gTTS(text=text, lang=locate)
            # Lưu file âm thanh trong một luồng riêng biệt
            await asyncio.to_thread(tts.save, audio_file_path)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        return FileResponse(
            audio_file_path, media_type="audio/wav", filename="output.wav"
        )
    elif file.filename.endswith(".docx"):
        # Đọc nội dung từ tệp Word
        content = await file.read()
        doc = Document(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")
        try:
            tts = gTTS(text=text, lang=locate)
            # Lưu file âm thanh trong một luồng riêng biệt
            await asyncio.to_thread(tts.save, audio_file_path)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        return FileResponse(
            audio_file_path, media_type="audio/wav", filename="output.wav"
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid file format.")


# API TFTS và Infer
@app.post("/text-file-to-speech-and-infer/")
async def text_file_to_speech_and_infer(
    file: UploadFile = File(...),
    locate: str = Form("vi"),
    model_id: int = Form(...),
    db: Session = Depends(get_db),
):
    if file.filename.endswith(".txt"):
        # Đọc nội dung từ tệp văn bản
        os.makedirs(section_storage_path, exist_ok=True)
        os.makedirs(train_model_path, exist_ok=True)

        content = await file.read()
        text = content.decode("utf-8")  # Giải mã nội dung

        # Lấy model_path và config_path từ cơ sở dữ liệu
        model_info = db.query(Model).filter(Model.id_model == model_id).first()
        if not model_info:
            raise HTTPException(status_code=400, detail="Model not found")

        model_path = BASE_DIR / model_info.model_path
        config_path = BASE_DIR / model_info.config_path
        cluster_model_path = BASE_DIR / model_info.cluster_model_path
        # Tạo tên file âm thanh ngẫu nhiên
        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

        # Chuyển đổi văn bản thành giọng nói
        try:
            tts = gTTS(text=text, lang=locate)
            # Lưu file âm thanh trong một luồng riêng biệt
            await asyncio.to_thread(tts.save, audio_file_path)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        output_audio_path = os.path.join(
            section_storage_path, f"{section_id}_processed.wav"
        )

        # Gọi hàm infer với file âm thanh vừa tạo ra
        try:
            await asyncio.to_thread(
                infer,
                input_path=audio_file_path,
                output_path=Path(output_audio_path),
                model_path=Path(model_path),
                config_path=Path(config_path),
                speaker=0,
                cluster_model_path=(
                    None
                    if model_info.cluster_model_path is None
                    else Path(cluster_model_path)
                ),
                transpose=0,
                auto_predict_f0=True,
                cluster_infer_ratio=0.0,
                noise_scale=0.4,
                f0_method="dio",
                db_thresh=-40,
                pad_seconds=0.5,
                chunk_seconds=0.5,
                absolute_thresh=False,
                max_chunk_seconds=40,
            )
            # print(f"Audio đã được xử lý và lưu tại: {output_audio_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error processing audio: {str(e)}"
            )

        return FileResponse(
            output_audio_path, media_type="audio/wav", filename="output.wav"
        )
    elif file.filename.endswith(".docx"):
        os.makedirs(section_storage_path, exist_ok=True)
        os.makedirs(train_model_path, exist_ok=True)
        # Đọc nội dung từ tệp Word
        content = await file.read()
        doc = Document(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
        # Lấy model_path và config_path từ cơ sở dữ liệu
        model_info = db.query(Model).filter(Model.id_model == model_id).first()
        if not model_info:
            raise HTTPException(status_code=400, detail="Model not found")

        model_path = BASE_DIR / model_info.model_path
        config_path = BASE_DIR / model_info.config_path
        cluster_model_path = BASE_DIR / model_info.cluster_model_path

        # Tạo tên file âm thanh ngẫu nhiên
        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

        # Chuyển đổi văn bản thành giọng nói
        try:
            tts = gTTS(text=text, lang=locate)
            # Lưu file âm thanh trong một luồng riêng biệt
            await asyncio.to_thread(tts.save, audio_file_path)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        output_audio_path = os.path.join(
            section_storage_path, f"{section_id}_processed.wav"
        )

        # Gọi hàm infer với file âm thanh vừa tạo ra
        try:
            await asyncio.to_thread(
                infer,
                input_path=audio_file_path,
                output_path=Path(output_audio_path),
                model_path=Path(model_path),
                config_path=Path(config_path),
                speaker=0,
                cluster_model_path=Path(cluster_model_path),
                transpose=0,
                auto_predict_f0=True,
                cluster_infer_ratio=0.0,
                noise_scale=0.4,
                f0_method="dio",
                db_thresh=-40,
                pad_seconds=0.5,
                chunk_seconds=0.5,
                absolute_thresh=False,
                max_chunk_seconds=40,
            )
            # print(f"Audio đã được xử lý và lưu tại: {output_audio_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error processing audio: {str(e)}"
            )

        return FileResponse(
            output_audio_path, media_type="audio/wav", filename="output.wav"
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid file format.")


@app.post("/text-to-speech-and-infer/")
async def text_to_speech_and_process(
    request: TextToSpeechAndInferRequest, db: Session = Depends(get_db)
):
    # Kiểm tra thông tin mô hình trong cơ sở dữ liệu
    model_info = db.query(Model).filter(Model.id_model == request.model_id).first()
    if not model_info:
        raise HTTPException(status_code=400, detail="Model not found")
    os.makedirs(section_storage_path, exist_ok=True)
    os.makedirs(train_model_path, exist_ok=True)
    # Lấy thông tin đường dẫn mô hình và cấu hình từ cơ sở dữ liệu
    model_path = BASE_DIR / model_info.model_path
    config_path = BASE_DIR / model_info.config_path
    cluster_model_path = BASE_DIR / model_info.cluster_model_path
    # Tạo tên file ngẫu nhiên
    section_id = str(uuid.uuid4())
    section_cloned_file_path = os.path.join(section_storage_path, section_id + ".wav")

    # Sử dụng gTTS để chuyển đổi văn bản thành giọng nói
    try:
        tts = gTTS(text=request.text, lang=request.locate)
        # Lưu file âm thanh trong một luồng riêng biệt
        await asyncio.to_thread(tts.save, section_cloned_file_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating audio: {str(e)}")

    # Gọi hàm infer để xử lý âm thanh (sử dụng asyncio.to_thread nếu infer là đồng bộ)
    output_audio_path = os.path.join(
        section_storage_path, f"{section_id}_processed.wav"
    )
    try:
        await asyncio.to_thread(
            infer,
            input_path=section_cloned_file_path,
            output_path=Path(output_audio_path),
            model_path=Path(model_path),
            config_path=Path(config_path),
            speaker=0,
            cluster_model_path=Path(cluster_model_path),
            transpose=0,
            auto_predict_f0=True,
            cluster_infer_ratio=0.0,
            noise_scale=0.4,
            f0_method="dio",
            db_thresh=-40,
            pad_seconds=0.5,
            chunk_seconds=0.5,
            absolute_thresh=False,
            max_chunk_seconds=40,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing audio: {str(e)}")

    return FileResponse(
        output_audio_path, media_type="audio/wav", filename="output.wav"
    )


@app.post("/infer-audio/", response_class=FileResponse)
async def upload_and_infer(
    file: UploadFile = File(...),
    model_id: int = Form(...),
    db: Session = Depends(get_db),
):

    # Kiểm tra thông tin mô hình trong cơ sở dữ liệu
    model_info = db.query(Model).filter(Model.id_model == model_id).first()
    if not model_info:
        raise HTTPException(status_code=400, detail="Model not found")
    os.makedirs(section_storage_path, exist_ok=True)
    os.makedirs(train_model_path, exist_ok=True)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())

    input_file_path = os.path.join(section_storage_path, f"{input_id}.wav")
    output_file_path = os.path.join(section_storage_path, f"{output_id}_processed.wav")

    cluster_model_path = BASE_DIR / model_info.cluster_model_path
    # Lưu tệp âm thanh được tải lên bằng aiofiles
    async with aiofiles.open(input_file_path, "wb") as buffer:
        content = await file.read()
        await buffer.write(content)

    # Gọi hàm infer để xử lý âm thanh (cần đảm bảo infer là hàm async hoặc sử dụng asyncio.to_thread nếu không đồng bộ)
    await asyncio.to_thread(
        infer,
        input_path=input_file_path,  # Sử dụng tệp đã lưu
        output_path=Path(output_file_path),
        model_path=Path(BASE_DIR / model_info.model_path),
        config_path=Path(BASE_DIR / model_info.config_path),
        speaker=0,
        cluster_model_path=Path(cluster_model_path),
        transpose=0,
        auto_predict_f0=True,
        cluster_infer_ratio=0.0,
        noise_scale=0.4,
        f0_method="dio",
        db_thresh=-40,
        pad_seconds=0.5,
        chunk_seconds=0.5,
        absolute_thresh=False,
        max_chunk_seconds=40,
    )

    print(f"Audio đã được xử lý và lưu tại: {output_file_path}")
    return FileResponse(output_file_path, media_type="audio/wav", filename="output.wav")


@app.post("/train-model/")
async def process_audio(
    name: str = Form(...), file: UploadFile = File(...), f0_method: str = Form(...)
):
    # Tạo thư mục tạm để lưu file âm thanh
    input_dir = BASE_DIR / "audio_data"
    output_dir = BASE_DIR / f"trainmodel/{name}"
    input_train = BASE_DIR / f"trainmodel/{name}/dataset/44k"
    output_split = BASE_DIR / f"trainmodel/{name}/dataset_raw"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    if not os.path.exists(input_dir):
        os.makedirs(input_dir, exist_ok=True)

    temp_file_path = Path(input_dir) / file.filename

    # Lưu tệp âm thanh vào thư mục
    if not os.path.exists(input_dir):
        os.makedirs(input_dir)

    await save_uploaded_file(file.file, temp_file_path)

    # Chạy lần lượt các hàm xử lý
    try:
        # Bước 1: Pre-split
        await asyncio.to_thread(
            pre_split,
            input_dir=input_dir,
            output_dir=os.path.join(output_dir, f"dataset_raw/{name}"),
            sr=22050,
        )
        # Bước 2: Process Audio Files
        await asyncio.to_thread(
            process_audio_files,
            input_dir=output_split,
            output_dir=input_train,
            sampling_rate=16000,
        )

        # Bước 3: Pre-config
        config_dir = os.path.join(output_dir, "configs/44k/config.json")
        filelist_dir = os.path.join(output_dir, "filelists/44k")
        config_type = "so-vits-svc-4.0v1"
        await asyncio.to_thread(
            pre_config,
            input_dir=input_train,
            filelist_path=filelist_dir,
            config_path=config_dir,
            config_type=config_type,
        )
        # Bước 4: Pre-hubert
        await asyncio.to_thread(
            pre_hubert,
            input_dir=input_train,
            config_path=config_dir,
            n_jobs=None,
            force_rebuild=True,
            f0_method=f0_method,
        )

        # Bước 5: Train model
        # model_dir = os.path.join(output_dir, "logs/44k")
        # await asyncio.to_thread(
        #     train,
        #     config_path=config_dir,
        #     model_path=model_dir,
        #     tensorboard=False,
        #     reset_optimizer=False,
        # )

        return {"message": "Audio processing completed successfully!"}

    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
