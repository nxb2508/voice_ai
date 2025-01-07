from fastapi import (
    FastAPI,
    UploadFile,
    File,
    HTTPException,
    Form,
    Response,
    Depends,
    Query,
    status,
    BackgroundTasks,
)
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import firebase_admin
from firebase_admin import credentials, firestore
from so_vits_svc_fork.inference.main import infer
from so_vits_svc_fork.preprocessing.preprocess_split import preprocess_split
from so_vits_svc_fork.preprocessing.preprocess_resample import preprocess_resample
from so_vits_svc_fork.preprocessing.preprocess_flist_config import (
    CONFIG_TEMPLATE_DIR,
    preprocess_config,
)
from so_vits_svc_fork.preprocessing.preprocess_hubert_f0 import preprocess_hubert_f0
from so_vits_svc_fork.train import train
import uvicorn
from gtts import gTTS
from pathlib import Path
import tempfile
from pydantic import BaseModel
from typing import List, Optional, Literal
import os
import uuid
from docx import Document
import io
import logging
from logging import getLogger
import shutil
import asyncio
import aiofiles
import re
from concurrent.futures import ThreadPoolExecutor
import json
import zipfile
import requests
from datetime import datetime


BASE_DIR = Path(__file__).resolve().parent


cred = credentials.Certificate(BASE_DIR / "APIkey.json")
firebase_admin.initialize_app(cred)
db_firestore = firestore.client()


section_storage_path = BASE_DIR / "files"
train_model_path = BASE_DIR / "trainmodel"

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

executor = ThreadPoolExecutor(max_workers=20)
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ModelResponse(BaseModel):
    id_model: str = None
    name_model: str
    model_name: str
    model_path: str
    config_path: str
    cluster_model_path: str
    category: str
    user_id: str
    created_at: str
    train_at: str

    class Config:
        from_attributes = True
        protected_namespaces = ()


class ModelCreate(BaseModel):
    name_model: str
    model_name: str
    model_path: str
    config_path: str
    cluster_model_path: str
    category: str
    user_id: str
    created_at: str
    train_at: str


class UpdateModelRequest(BaseModel):
    model_name: Optional[str]
    model_path: Optional[str]
    config_path: Optional[str]
    cluster_model_path: Optional[str]
    category: Optional[str]


class TextToSpeechRequest(BaseModel):
    text: str
    locate: str


class TextToSpeechAndInferRequest(BaseModel):
    text: str
    locate: str = "en"
    model_id: str


def get_current_time():
    return datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")


def check_directory_exists(folder_path: Path, subdir_name: str) -> bool:

    target_dir = folder_path / "trainmodel" / subdir_name
    return target_dir.exists() and target_dir.is_dir()


def get_latest_model_path(log_dir):
    max_epoch = -1
    model_path = None

    for filename in os.listdir(log_dir):

        match = re.match(r"G_(\d+)\.pth", filename)
        if match:
            epoch = int(match.group(1))

            if epoch > max_epoch:
                max_epoch = epoch
                model_path = Path(log_dir) / filename

    return str(model_path).replace("\\", "/") if model_path else None


async def save_uploaded_file(uploaded_file, destination: Path):
    try:
        with destination.open("wb") as dest_file:
            while True:
                chunk = await asyncio.to_thread(uploaded_file.read, 1024)
                if not chunk:
                    break
                dest_file.write(chunk)
    except Exception as e:
        raise RuntimeError(f"Error saving uploaded file: {e}")


def update_config(file_path, epochs_number):

    with open(file_path, "r") as file:
        config = json.load(file)

    config["train"]["epochs"] = epochs_number
    config["train"]["batch_size"] = 4
    config["train"]["log_interval"] = 200
    config["train"]["eval_interval"] = 400
    config["train"]["learning_rate"] = 0.0001

    with open(file_path, "w") as file:
        json.dump(config, file, indent=4)
    print("update thành công")


def pre_split(
    input_dir: str,
    output_dir: str,
    sr: int,
    max_length: float = 10.0,
    top_db: int = 30,
    frame_seconds: float = 1,
    hop_seconds: float = 0.3,
    n_jobs: int = -1,
) -> None:

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    preprocess_split(
        input_dir=input_dir,
        output_dir=output_dir,
        sr=sr,
        max_length=max_length,
        top_db=top_db,
        frame_seconds=frame_seconds,
        hop_seconds=hop_seconds,
        n_jobs=n_jobs,
    )


def pre_resample(
    input_dir: str,
    output_dir: str,
    sampling_rate: int = 44100,
    top_db: int = 30,
    frame_seconds: float = 1,
    hop_seconds: float = 0.3,
    n_jobs: int = -1,
) -> None:
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    output_path.mkdir(parents=True, exist_ok=True)

    preprocess_resample(
        input_dir=input_path,
        output_dir=output_path,
        sampling_rate=sampling_rate,
        n_jobs=n_jobs,
        top_db=top_db,
        frame_seconds=frame_seconds,
        hop_seconds=hop_seconds,
    )


def pre_config(
    input_dir: Path,
    filelist_path: Path,
    config_path: Path,
    config_type: str,
):

    input_dir = Path(input_dir)
    filelist_path = Path(filelist_path)
    config_path = Path(config_path)
    preprocess_config(
        input_dir=input_dir,
        train_list_path=filelist_path / "train.txt",
        val_list_path=filelist_path / "val.txt",
        test_list_path=filelist_path / "test.txt",
        config_path=config_path,
        config_name=config_type,
    )

    print("Configuration processing completed.")


def pre_hubert(
    input_dir: Path,
    config_path: Path,
    n_jobs: bool,
    force_rebuild: bool,
    f0_method: Literal["crepe", "crepe-tiny", "parselmouth", "dio", "harvest"],
) -> None:
    from so_vits_svc_fork.preprocessing.preprocess_hubert_f0 import preprocess_hubert_f0

    input_dir = Path(input_dir)
    config_path = Path(config_path)
    preprocess_hubert_f0(
        input_dir=input_dir,
        config_path=config_path,
        n_jobs=n_jobs,
        force_rebuild=force_rebuild,
        f0_method=f0_method,
    )


def train(
    config_path: Path,
    model_path: Path,
    tensorboard: bool = False,
    reset_optimizer: bool = False,
):
    config_path = Path(config_path)
    model_path = Path(model_path)

    if tensorboard:
        import webbrowser
        from tensorboard import program

        getLogger("tensorboard").setLevel(30)
        tb = program.TensorBoard()
        tb.configure(argv=[None, "--logdir", model_path.as_posix()])
        url = tb.launch()
        webbrowser.open(url)

    from so_vits_svc_fork.train import train

    train(
        config_path=config_path, model_path=model_path, reset_optimizer=reset_optimizer
    )


def train_and_save_model(
    name,
    file_name,
    input_dir,
    output_dir,
    input_train,
    output_split,
    f0_method,
    epochs_number,
    user_id,
    trainAt,
):
    try:

        pre_split(
            input_dir=input_dir,
            output_dir=os.path.join(output_dir, f"dataset_raw/{name}"),
            sr=44100,
        )
        pre_resample(
            input_dir=output_split, output_dir=input_train, sampling_rate=44100
        )

        config_dir = os.path.join(output_dir, "configs/44k/config.json")
        filelist_dir = os.path.join(output_dir, "filelists/44k")
        config_type = "so-vits-svc-4.0v1"
        pre_config(
            input_dir=input_train,
            filelist_path=filelist_dir,
            config_path=config_dir,
            config_type=config_type,
        )

        update_config(config_dir, epochs_number)
        pre_hubert(
            input_dir=input_train,
            config_path=config_dir,
            n_jobs=None,
            force_rebuild=True,
            f0_method=f0_method,
        )

        model_dir = os.path.join(output_dir, "logs/44k")
        train(
            config_path=config_dir,
            model_path=model_dir,
            tensorboard=False,
            reset_optimizer=False,
        )

        latest_model_path = get_latest_model_path(model_dir)
        config_path_relative = Path(config_dir).relative_to(BASE_DIR).as_posix()
        latest_model_path_relative = (
            Path(latest_model_path).relative_to(BASE_DIR).as_posix()
        )

        created_at = get_current_time()
        model_id = str(uuid.uuid4())
        model_ref = db_firestore.collection("models").document(model_id)
        model_data = {
            "name_model": name,
            "model_name": file_name,
            "model_path": latest_model_path_relative,
            "config_path": config_path_relative,
            "cluster_model_path": "None",
            "category": "1",
            "user_id": user_id,
            "created_at": created_at,
            "train_at": trainAt,
        }
        model_ref.set(model_data)

    except Exception as e:
        print(f"Error during training: {e}")


# Xóa một model theo ID
@app.delete("/models/{model_id}")
async def delete_model(model_id: str):
    try:

        output_dir = BASE_DIR / f"trainmodel"
        model_doc = db_firestore.collection("models").document(str(model_id)).get()
        if not model_doc.exists:
            raise HTTPException(status_code=404, detail="Model không tồn tại.")

        model_data = model_doc.to_dict()
        model_name = model_data.get("model_name")

        model_folder_path = os.path.join(output_dir, model_name)

        if os.path.exists(model_folder_path):
            shutil.rmtree(model_folder_path)
            logger.info(f"Đã xóa thư mục: {model_folder_path}")
        else:
            logger.warning(f"Thư mục không tồn tại: {model_folder_path}")

        db_firestore.collection("models").document(str(model_id)).delete()
        return {"message": "Xóa model thành công."}

    except Exception as e:
        logger.error(f"Lỗi khi xóa model: {e}")
        raise HTTPException(status_code=500, detail="Không thể xóa model.")


@app.delete("/models/{model_id}/cleanup")
async def cleanup_model_files(model_id: str):
    try:
        output_dir = BASE_DIR / f"trainmodel"

        model_doc = db_firestore.collection("models").document(str(model_id)).get()
        if not model_doc.exists:
            raise HTTPException(status_code=404, detail="Model không tồn tại.")

        model_data = model_doc.to_dict()
        model_name = model_data.get("model_name")

        folder_path = os.path.join(output_dir, model_name, "logs", "44k")
        if not os.path.exists(folder_path):
            raise HTTPException(status_code=404, detail="Đường dẫn không tồn tại.")

        files = [
            f
            for f in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, f))
        ]
        max_d_file = None
        max_g_file = None
        max_d_version = -1
        max_g_version = -1

        for file in files:
            if file.startswith("D_") and file.endswith(".pth"):
                version = int(file.split("_")[1].split(".")[0])
                if version > max_d_version:
                    max_d_version = version
                    max_d_file = file
            elif file.startswith("G_") and file.endswith(".pth"):
                version = int(file.split("_")[1].split(".")[0])
                if version > max_g_version:
                    max_g_version = version
                    max_g_file = file

        for file in files:
            file_path = os.path.join(folder_path, file)
            if file.startswith("D_") and file != max_d_file:
                os.remove(file_path)
                print(f"Đã xóa: {file_path}")
            elif file.startswith("G_") and file != max_g_file:
                os.remove(file_path)
                print(f"Đã xóa: {file_path}")

        return {
            "message": "Dọn dẹp file thành công.",
            "remaining_files": [max_d_file, max_g_file],
        }

    except Exception as e:
        print(f"Lỗi khi dọn dẹp tệp: {e}")
        raise HTTPException(status_code=500, detail=f"Lỗi khi dọn dẹp file: {str(e)}")


# API TTS
@app.post("/text-to-speech/")
async def text_to_speech(request: TextToSpeechRequest):
    os.makedirs(section_storage_path, exist_ok=True)
    section_id = str(uuid.uuid4())
    audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

    try:
        tts = gTTS(text=text, lang=locate)
        await asyncio.to_thread(tts.save, audio_file_path)
        # print(f"File âm thanh đã được lưu tại: {audio_file_path}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating audio: {str(e)}")


@app.post("/text-file-to-speech/")
async def text_file_to_speech(
    file: UploadFile = File(...),
    locate: str = Form("vi"),
):
    if file.filename.endswith(".txt"):

        content = await file.read()
        text = content.decode("utf-8")
        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

        try:
            tts = gTTS(text=text, lang=locate)
            await asyncio.to_thread(tts.save, audio_file_path)
            # print(f"File âm thanh đã được lưu tại: {audio_file_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        return FileResponse(
            audio_file_path, media_type="audio/wav", filename="output.wav"
        )
    elif file.filename.endswith(".docx"):

        content = await file.read()
        doc = Document(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")
        try:
            tts = gTTS(text=text, lang=locate)
            await asyncio.to_thread(tts.save, audio_file_path)
            # print(f"File âm thanh đã được lưu tại: {audio_file_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        return FileResponse(
            audio_file_path, media_type="audio/wav", filename="output.wav"
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid file format.")


# API TFTS và Infer
@app.post("/text-file-to-speech-and-infer/")
async def text_file_to_speech_and_infer(
    file: UploadFile = File(...),
    locate: str = Form(...),
    model_id: str = Form(...),
):
    if file.filename.endswith(".txt"):

        os.makedirs(section_storage_path, exist_ok=True)
        os.makedirs(train_model_path, exist_ok=True)

        content = await file.read()
        text = content.decode("utf-8")

        try:
            model_doc = db_firestore.collection("models").document(str(model_id)).get()
            if not model_doc.exists:
                raise HTTPException(status_code=400, detail="Model not found")

            model_info = model_doc.to_dict()
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Lỗi khi lấy thông tin model: {e}"
            )

        model_path = BASE_DIR / model_info["model_path"]
        config_path = BASE_DIR / model_info["config_path"]
        cluster_model_path = BASE_DIR / model_info["cluster_model_path"]

        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

        try:
            tts = gTTS(text=text, lang=locate)
            await asyncio.to_thread(tts.save, audio_file_path)
            # print(f"File âm thanh đã được lưu tại: {audio_file_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        output_audio_path = os.path.join(
            section_storage_path, f"{section_id}_processed.wav"
        )

        try:
            await asyncio.to_thread(
                infer,
                input_path=audio_file_path,
                output_path=Path(output_audio_path),
                model_path=Path(model_path),
                config_path=Path(config_path),
                speaker=0,
                cluster_model_path=Path(cluster_model_path),
                transpose=0,
                auto_predict_f0=True,
                cluster_infer_ratio=0.0,
                noise_scale=0.4,
                f0_method="dio",
                db_thresh=-40,
                pad_seconds=0.5,
                chunk_seconds=0.5,
                absolute_thresh=False,
                max_chunk_seconds=40,
            )
            # print(f"Audio đã được xử lý và lưu tại: {output_audio_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error processing audio: {str(e)}"
            )

        return FileResponse(
            output_audio_path, media_type="audio/wav", filename="output.wav"
        )
    elif file.filename.endswith(".docx"):
        os.makedirs(section_storage_path, exist_ok=True)
        os.makedirs(train_model_path, exist_ok=True)

        content = await file.read()
        doc = Document(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])

        try:
            model_doc = db_firestore.collection("models").document(str(model_id)).get()
            if not model_doc.exists:
                raise HTTPException(status_code=400, detail="Model not found")

            model_info = model_doc.to_dict()
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Lỗi khi lấy thông tin model: {e}"
            )

        model_path = BASE_DIR / model_info["model_path"]
        config_path = BASE_DIR / model_info["config_path"]
        cluster_model_path = BASE_DIR / model_info["cluster_model_path"]

        section_id = str(uuid.uuid4())
        audio_file_path = os.path.join(section_storage_path, section_id + ".wav")

        try:
            tts = gTTS(text=text, lang=locate)
            await asyncio.to_thread(tts.save, audio_file_path)
            # print(f"File âm thanh đã được lưu tại: {audio_file_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error generating audio: {str(e)}"
            )

        output_audio_path = os.path.join(
            section_storage_path, f"{section_id}_processed.wav"
        )

        try:
            await asyncio.to_thread(
                infer,
                input_path=audio_file_path,
                output_path=Path(output_audio_path),
                model_path=Path(model_path),
                config_path=Path(config_path),
                speaker=0,
                cluster_model_path=Path(cluster_model_path),
                transpose=0,
                auto_predict_f0=True,
                cluster_infer_ratio=0.0,
                noise_scale=0.4,
                f0_method="dio",
                db_thresh=-40,
                pad_seconds=0.5,
                chunk_seconds=0.5,
                absolute_thresh=False,
                max_chunk_seconds=40,
            )
            # print(f"Audio đã được xử lý và lưu tại: {output_audio_path}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error processing audio: {str(e)}"
            )

        return FileResponse(
            output_audio_path, media_type="audio/wav", filename="output.wav"
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid file format.")


@app.post("/text-to-speech-and-infer/")
async def text_to_speech_and_process(request: TextToSpeechAndInferRequest):
    model_id = request.model_id
    text = request.text
    locate = request.locate

    os.makedirs(section_storage_path, exist_ok=True)
    os.makedirs(train_model_path, exist_ok=True)

    try:
        model_doc = db_firestore.collection("models").document(str(model_id)).get()
        if not model_doc.exists:
            raise HTTPException(status_code=400, detail="Model not found")

        model_info = model_doc.to_dict()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi khi lấy thông tin model: {e}")

    model_path = BASE_DIR / model_info["model_path"]
    config_path = BASE_DIR / model_info["config_path"]
    cluster_model_path = BASE_DIR / model_info["cluster_model_path"]

    section_id = str(uuid.uuid4())
    section_cloned_file_path = os.path.join(section_storage_path, section_id + ".wav")

    try:
        tts = gTTS(text=text, lang=locate)
        await asyncio.to_thread(tts.save, section_cloned_file_path)
        # print(f"File âm thanh đã được lưu tại: {audio_file_path}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating audio: {str(e)}")

    output_audio_path = os.path.join(
        section_storage_path, f"{section_id}_processed.wav"
    )
    try:
        await asyncio.to_thread(
            infer,
            input_path=section_cloned_file_path,
            output_path=Path(output_audio_path),
            model_path=Path(model_path),
            config_path=Path(config_path),
            speaker=0,
            cluster_model_path=Path(cluster_model_path),
            transpose=0,
            auto_predict_f0=True,
            cluster_infer_ratio=0.0,
            noise_scale=0.4,
            f0_method="dio",
            db_thresh=-40,
            pad_seconds=0.5,
            chunk_seconds=0.5,
            absolute_thresh=False,
            max_chunk_seconds=40,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing audio: {str(e)}")

    return FileResponse(
        output_audio_path, media_type="audio/wav", filename="output.wav"
    )


@app.post("/infer-audio/", response_class=FileResponse)
async def upload_and_infer(
    file: UploadFile = File(...),
    model_id: str = Form(...),
):

    try:
        model_doc = db_firestore.collection("models").document(str(model_id)).get()
        if not model_doc.exists:
            raise HTTPException(status_code=400, detail="Model not found")

        model_info = model_doc.to_dict()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi khi lấy thông tin model: {e}")

    os.makedirs(section_storage_path, exist_ok=True)
    os.makedirs(train_model_path, exist_ok=True)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())

    input_file_path = os.path.join(section_storage_path, f"{input_id}.wav")
    output_file_path = os.path.join(section_storage_path, f"{output_id}_processed.wav")

    cluster_model_path = BASE_DIR / model_info["cluster_model_path"]

    async with aiofiles.open(input_file_path, "wb") as buffer:
        content = await file.read()
        await buffer.write(content)

    await asyncio.to_thread(
        infer,
        input_path=input_file_path,
        output_path=Path(output_file_path),
        model_path=Path(BASE_DIR / model_info["model_path"]),
        config_path=Path(BASE_DIR / model_info["config_path"]),
        speaker=0,
        cluster_model_path=Path(cluster_model_path),
        transpose=0,
        auto_predict_f0=True,
        cluster_infer_ratio=0.0,
        noise_scale=0.4,
        f0_method="dio",
        db_thresh=-40,
        pad_seconds=0.5,
        chunk_seconds=0.5,
        absolute_thresh=False,
        max_chunk_seconds=40,
    )

    print(f"Audio đã được xử lý và lưu tại: {output_file_path}")
    return FileResponse(output_file_path, media_type="audio/wav", filename="output.wav")


# @app.post("/train-model/")
# async def process_audio(
#     name: str = Form(...),
#     file: UploadFile = File(...),
#     f0_method: str = Form(...),
#     epochs_number: str = Form(...),
#     user_id: str = Form(...),
#     trainAt: str = Form(...),
# ):

#     suid = str(uuid.uuid4())[:8]
#     file_name = file_name = name + "_" + epochs_number + "_" + suid
#     input_dir = BASE_DIR / f"audio_data/{file_name}"
#     output_dir = BASE_DIR / f"trainmodel/{file_name}"
#     input_train = BASE_DIR / f"trainmodel/{file_name}/dataset/44k"
#     output_split = BASE_DIR / f"trainmodel/{file_name}/dataset_raw"
#     if not os.path.exists(output_dir):
#         os.makedirs(output_dir)
#     if not os.path.exists(input_dir):
#         os.makedirs(input_dir, exist_ok=True)

#     temp_file_path = Path(input_dir) / file.filename

#     if not os.path.exists(input_dir):
#         os.makedirs(input_dir)

#     await save_uploaded_file(file.file, temp_file_path)

#     try:

#         await asyncio.to_thread(
#             pre_split,
#             input_dir=input_dir,
#             output_dir=os.path.join(output_dir, f"dataset_raw/{name}"),
#             sr=44100,
#         )

#         await asyncio.to_thread(
#             pre_resample,
#             input_dir=output_split,
#             output_dir=input_train,
#             sampling_rate=44100,
#         )

#         config_dir = os.path.join(output_dir, "configs/44k/config.json")
#         filelist_dir = os.path.join(output_dir, "filelists/44k")
#         config_type = "so-vits-svc-4.0v1"
#         await asyncio.to_thread(
#             pre_config,
#             input_dir=input_train,
#             filelist_path=filelist_dir,
#             config_path=config_dir,
#             config_type=config_type,
#         )
#         epochs_number = int(epochs_number)
#         update_config(config_dir, epochs_number)

#         await asyncio.to_thread(
#             pre_hubert,
#             input_dir=input_train,
#             config_path=config_dir,
#             n_jobs=None,
#             force_rebuild=True,
#             f0_method=f0_method,
#         )

#         model_dir = os.path.join(output_dir, "logs/44k")
#         await asyncio.to_thread(
#             train,
#             config_path=config_dir,
#             model_path=model_dir,
#             tensorboard=False,
#             reset_optimizer=False,
#         )
#         latest_model_path = get_latest_model_path(model_dir)
#         config_path_relative = Path(config_dir).relative_to(BASE_DIR).as_posix()
#         latest_model_path_relative = (
#             Path(latest_model_path).relative_to(BASE_DIR).as_posix()
#         )
#         created_at = get_current_time
#         model_id = str(uuid.uuid4())
#         model_ref = db_firestore.collection("models").document(model_id)
#         model_data = {
#             "name_model": name,
#             "model_name": file_name,
#             "model_path": latest_model_path_relative,
#             "config_path": config_path_relative,
#             "cluster_model_path": "None",
#             "category": "1",
#             "user_id": user_id,
#             "created_at": created_at,
#             "train_at": trainAt,
#         }
#         model_ref.set(model_data)
#         return {
#             "message": "Audio processing completed successfully!",
#         }
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


@app.post("/train-model-file-zip/")
async def process_audio_zip(
    name: str = Form(...),
    file: UploadFile = File(...),
    f0_method: str = Form(...),
    epochs_number: str = Form(...),
    user_id: str = Form(...),
    trainAt: str = Form(...),
):

    suid = str(uuid.uuid4())[:8]

    file_name = name + "_" + epochs_number + "_" + suid
    input_dir = BASE_DIR / f"audio_data/{file_name}"
    output_dir = BASE_DIR / f"trainmodel/{file_name}"
    input_train = BASE_DIR / f"trainmodel/{file_name}/dataset/44k"
    output_split = BASE_DIR / f"trainmodel/{file_name}/dataset_raw"
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(input_dir, exist_ok=True)

    temp_file_path = Path(input_dir) / file.filename

    await save_uploaded_file(file.file, temp_file_path)

    try:

        if file.filename.endswith(".zip"):
            with zipfile.ZipFile(temp_file_path, "r") as zip_ref:
                zip_ref.extractall(output_split)

            await asyncio.to_thread(
                pre_resample,
                input_dir=output_split,
                output_dir=input_train,
                sampling_rate=44100,
            )
            config_dir = os.path.join(output_dir, "configs/44k/config.json")
            filelist_dir = os.path.join(output_dir, "filelists/44k")
            config_type = "so-vits-svc-4.0v1"
            await asyncio.to_thread(
                pre_config,
                input_dir=input_train,
                filelist_path=filelist_dir,
                config_path=config_dir,
                config_type=config_type,
            )
            epochs_number = int(epochs_number)
            update_config(config_dir, epochs_number)
            await asyncio.to_thread(
                pre_hubert,
                input_dir=input_train,
                config_path=config_dir,
                n_jobs=None,
                force_rebuild=True,
                f0_method=f0_method,
            )
            model_dir = os.path.join(output_dir, "logs/44k")
            await asyncio.to_thread(
                train,
                config_path=config_dir,
                model_path=model_dir,
                tensorboard=False,
                reset_optimizer=False,
            )
            latest_model_path = get_latest_model_path(model_dir)
            config_path_relative = Path(config_dir).relative_to(BASE_DIR).as_posix()
            latest_model_path_relative = (
                Path(latest_model_path).relative_to(BASE_DIR).as_posix()
            )
            created_at = get_current_time
            model_id = str(uuid.uuid4())
            model_ref = db_firestore.collection("models").document(model_id)
            model_data = {
                "name_model": name,
                "model_name": file_name,
                "model_path": latest_model_path_relative,
                "config_path": config_path_relative,
                "cluster_model_path": "None",
                "category": "1",
                "user_id": user_id,
                "created_at": created_at,
                "train_at": trainAt,
            }
            model_ref.set(model_data)
            return {
                "message": "Audio processing completed successfully!",
            }
        else:
            await asyncio.to_thread(
                pre_split,
                input_dir=input_dir,
                output_dir=os.path.join(output_dir, f"dataset_raw/{name}"),
                sr=44100,
            )
            await asyncio.to_thread(
                pre_resample,
                input_dir=output_split,
                output_dir=input_train,
                sampling_rate=44100,
            )

            config_dir = os.path.join(output_dir, "configs/44k/config.json")
            filelist_dir = os.path.join(output_dir, "filelists/44k")
            config_type = "so-vits-svc-4.0v1"
            await asyncio.to_thread(
                pre_config,
                input_dir=input_train,
                filelist_path=filelist_dir,
                config_path=config_dir,
                config_type=config_type,
            )
            epochs_number = int(epochs_number)
            update_config(config_dir, epochs_number)
            await asyncio.to_thread(
                pre_hubert,
                input_dir=input_train,
                config_path=config_dir,
                n_jobs=None,
                force_rebuild=True,
                f0_method=f0_method,
            )
            model_dir = os.path.join(output_dir, "logs/44k")
            await asyncio.to_thread(
                train,
                config_path=config_dir,
                model_path=model_dir,
                tensorboard=False,
                reset_optimizer=False,
            )
            latest_model_path = get_latest_model_path(model_dir)
            config_path_relative = Path(config_dir).relative_to(BASE_DIR).as_posix()
            latest_model_path_relative = (
                Path(latest_model_path).relative_to(BASE_DIR).as_posix()
            )
            created_at = get_current_time
            model_id = str(uuid.uuid4())
            model_ref = db_firestore.collection("models").document(model_id)
            model_data = {
                "name_model": name,
                "model_name": file_name,
                "model_path": latest_model_path_relative,
                "config_path": config_path_relative,
                "cluster_model_path": "None",
                "category": "1",
                "user_id": user_id,
                "created_at": created_at,
                "train_at": trainAt,
            }
            model_ref.set(model_data)
            return {
                "message": "Audio processing completed successfully!",
            }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/train-model/")
async def process_audio(
    background_tasks: BackgroundTasks,
    name: str = Form(...),
    file: UploadFile = File(...),
    f0_method: str = Form(...),
    epochs_number: str = Form(...),
    user_id: str = Form(...),
    trainAt: str = Form(...),
):
    suid = str(uuid.uuid4())[:8]
    file_name = name + "_" + epochs_number + "_" + suid
    input_dir = BASE_DIR / f"audio_data/{file_name}"
    output_dir = BASE_DIR / f"trainmodel/{file_name}"
    input_train = BASE_DIR / f"trainmodel/{file_name}/dataset/44k"
    output_split = BASE_DIR / f"trainmodel/{file_name}/dataset_raw"

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    if not os.path.exists(input_dir):
        os.makedirs(input_dir, exist_ok=True)

    temp_file_path = Path(input_dir) / file.filename
    await save_uploaded_file(file.file, temp_file_path)

    background_tasks.add_task(
        train_and_save_model,
        name,
        file_name,
        input_dir,
        output_dir,
        input_train,
        output_split,
        f0_method,
        int(epochs_number),
        user_id,
        trainAt,
    )

    return {"message": "Training started!"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
